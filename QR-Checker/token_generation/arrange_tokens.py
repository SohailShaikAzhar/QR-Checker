from docx import Document
from docx.shared import Inches
import os

# Paths
image_folder = r"PATH"  # Update this to your actual image folder
output_docx = r"PATH"

# Get all image files sorted
image_files = sorted([f for f in os.listdir(image_folder) if f.endswith(".png")])

# Create a Word document
doc = Document()

# Set document margins to 0.5 inches
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Function to create a 2-column table (8 images per page, 4 rows)
def add_images_to_table(doc, images):
    table = doc.add_table(rows=4, cols=2)  # 4 rows, 2 columns (8 images per page)
    table.autofit = False  # Fix cell sizes

    for i, image in enumerate(images):
        row = i // 2  # Determine row index (0-3)
        col = i % 2   # Determine column index (0-1)
        
        cell = table.cell(row, col)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        image_path = os.path.join(image_folder, image)
        run.add_picture(image_path, width=Inches(3.7), height=Inches(2.18))  # Set exact size

# Loop through images and insert 8 per page
for i in range(0, len(image_files), 8):
    add_images_to_table(doc, image_files[i:i+8])  # Add 8 images per table
    doc.add_page_break()  # Page break after every 8 images

# Save document
doc.save(output_docx)
print(f"✅ Document saved as: {output_docx}")
